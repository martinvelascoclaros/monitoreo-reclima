# -*- coding: utf-8 -*-
"""
Dashboard de monitoreo de calidad — RECLIMA (SCALL y Agrícola)
Publica los datos el administrador (carpeta data/ del repo); el equipo
consulta con enlace + contraseña. El tablero NO muestra datos personales
de los entrevistados (solicitud FAO): solo el ID (PRODUCTOS-ID ENCUESTA).

Banderas sustantivas (derivadas del análisis de calidad, jul 2026):
  Q   Registro no trazable — sin identificador (ambos)
  S   Composición del hogar incompleta (SCALL)
  X   Contradicción lógica (ambos módulos)
  A1a Sin cultivos capturados (Agrícola)
  A1b Roster productivo incompleto — menos de lo declarado (Agrícola)
  A2  Polígono pendiente de medir (Agrícola)
  A3  Medida del polígono muy diferente a la declarada (Agrícola)
"""

import io
import os
import re
import math
import unicodedata
from datetime import datetime

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

st.set_page_config(page_title="Monitoreo RECLIMA", page_icon="🌱",
                   layout="wide", initial_sidebar_state="collapsed")

# ----------------------------------------------------------------------------
# Estética
# ----------------------------------------------------------------------------
VERDE = "#2E7D46"
VERDE_OSC = "#1F5C2E"
AMBAR = "#E8A33D"
ROJO = "#D1495B"
GRIS = "#6b7280"

CSS = f"""
<style>
.stApp {{ background: #F7F9F6; }}
.block-container {{ padding-top: 2.2rem; max-width: 1300px; }}
h1, h2, h3 {{ color: {VERDE_OSC}; font-family: 'Segoe UI', system-ui, sans-serif; }}
[data-testid="stMetric"] {{
    background: #ffffff; border: 1px solid #e6ece6; border-radius: 14px;
    padding: 14px 16px; box-shadow: 0 1px 3px rgba(31,92,46,.06);
}}
[data-testid="stMetricLabel"] {{ color: {GRIS}; font-weight: 600; }}
[data-testid="stMetricValue"] {{ color: {VERDE_OSC}; }}
.stTabs [data-baseweb="tab-list"] {{ gap: 6px; }}
.stTabs [data-baseweb="tab"] {{
    background: #eef3ee; border-radius: 10px 10px 0 0; padding: 8px 18px;
    font-weight: 600;
}}
.stTabs [aria-selected="true"] {{ background: {VERDE}; color: #fff; }}
div[data-testid="stExpander"] {{ border-radius: 12px; border: 1px solid #e6ece6; }}
.hero {{
    background: linear-gradient(100deg, {VERDE_OSC}, {VERDE});
    color: #fff; padding: 20px 26px; border-radius: 16px; margin-bottom: 8px;
}}
.hero h1 {{ color: #fff; margin: 0; font-size: 1.7rem; }}
.hero p {{ color: #eafaef; margin: 4px 0 0; font-size: .95rem; }}
.pill {{ display:inline-block; padding:3px 10px; border-radius:999px;
    font-size:.78rem; font-weight:600; margin-right:6px; }}
.pill-red {{ background:#fdeaed; color:{ROJO}; }}
.pill-amber {{ background:#fdf3e3; color:#b9791f; }}
.pill-green {{ background:#e8f5ec; color:{VERDE_OSC}; }}
</style>
"""


# ----------------------------------------------------------------------------
# Acceso con contraseña
# ----------------------------------------------------------------------------
def check_password() -> bool:
    try:
        expected = st.secrets.get("APP_PASSWORD", "reclima2026")
    except Exception:
        expected = "reclima2026"
    if st.session_state.get("auth_ok"):
        return True
    st.markdown(CSS, unsafe_allow_html=True)
    st.title("🌱 Monitoreo RECLIMA")
    pwd = st.text_input("Contraseña", type="password")
    if pwd:
        if pwd == expected:
            st.session_state["auth_ok"] = True
            st.rerun()
        else:
            st.error("Contraseña incorrecta.")
    return False


# ----------------------------------------------------------------------------
# Utilidades
# ----------------------------------------------------------------------------
def _norm(s: str) -> str:
    s = unicodedata.normalize("NFKD", str(s))
    s = "".join(c for c in s if not unicodedata.combining(c))
    return s.lower()


def resolve(df: pd.DataFrame, key, exact: bool = False):
    keys = [key] if isinstance(key, str) else list(key)
    for k in keys:
        nk = _norm(k)
        for c in df.columns:
            if _norm(c) == nk:
                return c
    if exact:
        return None
    for k in keys:
        nk = _norm(k)
        for c in df.columns:
            if nk in _norm(c):
                return c
    return None


K_ENUM = ["Nombre del encuestador", "Nombre del enumerador"]
K_NOMBRE = ["nombre completo del productor",
            "nombre completo de la jefa o jefe", "nombre completo"]
K_ID = ["PRODUCTOS-ID ENCUESTA", "PRODUCTOS-ID", "ID ENCUESTA",
        "Identificador de encuestado"]


def col(df, key, exact=False):
    c = resolve(df, key, exact)
    if c is None:
        return pd.Series(np.nan, index=df.index)
    return df[c]


def num(s: pd.Series) -> pd.Series:
    return pd.to_numeric(s, errors="coerce")


def vacios(serie: pd.Series) -> pd.Series:
    return serie.isna() | serie.astype(str).str.strip().isin(["", "nan", "None", "<NA>"])


# Nombres posibles de una llave única creada por el usuario para unir las
# bases (DATA con los rosters). Si existe, tiene prioridad sobre la de Kobo.
K_LLAVE_DATA = ["INDEX", "ID_ENCUESTA", "ID ENCUESTA", "llave", "key_encuesta"]
K_LLAVE_ROSTER = ["INDEX", "ID_ENCUESTA", "ID ENCUESTA", "llave", "key_encuesta"]


def clave_encuesta_data(d):
    """Columna con la llave ÚNICA de encuesta en DATA. Prioriza una llave
    creada por el usuario (INDEX/ID_ENCUESTA); si no, usa el _id de Kobo
    (único). Evita _index, que se repite al mezclar versiones del formulario."""
    for k in K_LLAVE_DATA:
        c = resolve(d, k, exact=True)
        if c is not None and d[c].nunique(dropna=True) == d[c].notna().sum():
            return c
    return resolve(d, "_id", exact=True) or resolve(d, "_index", exact=True)


def clave_encuesta_roster(r):
    """Columna del roster que referencia a la encuesta padre. Prioriza la
    llave del usuario (INDEX/ID_ENCUESTA); si no, usa _submission__id de Kobo
    (== _id de DATA), globalmente único a diferencia de _parent_index."""
    for k in K_LLAVE_ROSTER:
        c = resolve(r, k, exact=True)
        if c is not None:
            return c
    return resolve(r, "_submission__id", exact=True)


def _kobo(d):
    kc = clave_encuesta_data(d)
    if kc is not None:
        s = d[kc]
        # La llave puede ser numérica (Kobo _id: 22180444.0 -> 22180444) o de
        # texto (columna INDEX creada por el usuario: "V3-R3"). No forzar entero.
        if pd.api.types.is_numeric_dtype(s):
            return s.astype("Int64").astype(str)
        return s.astype(str).str.strip()
    return pd.Series((d.index + 1).astype(str), index=d.index)


def id_encuesta(d: pd.DataFrame) -> pd.Series:
    # Si el usuario creó una llave propia (INDEX), esa es el identificador
    # visible del registro (única y siempre presente).
    for k in K_LLAVE_DATA:
        c = resolve(d, k, exact=True)
        if c is not None:
            return d[c].astype(str).str.strip()
    kobo = _kobo(d)
    id_c = resolve(d, K_ID)
    if id_c is not None:
        raw = d[id_c]
        txt = raw.astype(str).str.strip()
        sin = raw.isna() | txt.isin(["", "nan", "None", "<NA>"])
        return txt.mask(sin, "s/ID·k" + kobo).astype(str)
    return kobo


# --- Protección de datos personales ---
PATRON_SENSIBLE = re.compile(
    r"nombre|telefono|celular|correo|contacto|domicilio|direccion|geoloc"
    r"|latitud|longitud|latitude|longitude|altitude|precision|gps|poligono"
    r"|polygon|shape|coordenadas",
    re.I)


def es_sensible(nombre_col: str) -> bool:
    n = _norm(str(nombre_col))
    if "encuestador" in n or "supervisor" in n or "enumerador" in n:
        return False
    return bool(PATRON_SENSIBLE.search(n))


def quitar_sensibles(df: pd.DataFrame) -> pd.DataFrame:
    return df[[c for c in df.columns if not es_sensible(c)]]


@st.cache_data(show_spinner=False)
def load_book(file_bytes: bytes) -> dict:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return {name: xls.parse(name) for name in xls.sheet_names}


def pick_sheet(book: dict, *candidates):
    for cand in candidates:
        for name in book:
            if _norm(cand) == _norm(name):
                return book[name]
    for cand in candidates:
        for name in book:
            if _norm(cand) in _norm(name):
                return book[name]
    return None


def detect_module(book: dict):
    data = pick_sheet(book, "DATA")
    if data is None:
        data = max(book.values(), key=lambda d: d.shape[1])
    joined = " ".join(_norm(c) for c in data.columns)
    if "recoleccion de lluvia" in joined or "captacion de agua de lluvia" in joined or "scall" in joined:
        return "SCALL", data
    if "terrenos o parcelas" in joined or pick_sheet(book, "roster_parcela") is not None:
        return "AGRICOLA", data
    return None, data


# ----------------------------------------------------------------------------
# Geometría (polígonos de parcela)
# ----------------------------------------------------------------------------
M2_POR_MANZANA = 6989.0
# Factores de conversión a METROS CUADRADOS según la unidad que reporta el
# productor. La comparación de áreas siempre se hace en m².
def factor_a_m2(unidad):
    u = _norm(unidad)
    if "manzana" in u:
        return 6989.0            # 1 manzana = 6,989 m²
    if "hectarea" in u:
        return 10000.0           # 1 ha = 10,000 m²
    if "tarea" in u:
        return 437.0             # 1 tarea = 437 m² (1/16 de manzana)
    if "vara" in u:
        return 0.698737          # 1 vara² ≈ 0.6987 m²
    if "cuerda" in u:
        return 436.6             # aprox. (variable por región)
    if "metro" in u or u.strip() in ("m2", "m²"):
        return 1.0
    return None


def area_declarada_m2(area, unidad, especifique=None):
    """Devuelve (área en m², texto de la unidad). Si la unidad es 'Otra',
    usa el campo Especifique. Si no se puede convertir, devuelve (NaN, texto)."""
    a = pd.to_numeric(pd.Series([area]), errors="coerce").iloc[0]
    if pd.isna(a) or a <= 0:
        return np.nan, "—"
    fac = factor_a_m2(unidad)
    txt = str(unidad)
    if fac is None and especifique and str(especifique).strip() not in ("", "nan", "None"):
        fac = factor_a_m2(especifique)
        txt = str(especifique)
    if fac is None:
        return np.nan, f"{a:g} {unidad} (unidad no convertible)"
    return a * fac, f"{a:g} {txt}"


def parse_geoshape(s):
    pts = []
    for seg in str(s).split(";"):
        p = seg.strip().split()
        if len(p) >= 2:
            try:
                prec = float(p[3]) if len(p) > 3 else np.nan
                pts.append((float(p[0]), float(p[1]), prec))
            except ValueError:
                pass
    return pts


def area_poligono_m2(pts) -> float:
    if len(pts) < 3:
        return 0.0
    lat0 = sum(p[0] for p in pts) / len(pts)
    xs = [p[1] * 111320 * math.cos(math.radians(lat0)) for p in pts]
    ys = [p[0] * 110540 for p in pts]
    a = 0.0
    for i in range(len(xs)):
        j = (i + 1) % len(xs)
        a += xs[i] * ys[j] - xs[j] * ys[i]
    return abs(a) / 2


# ----------------------------------------------------------------------------
# BANDERAS SUSTANTIVAS
# ----------------------------------------------------------------------------
FLAG_DESC = {
    "Q Sin identificador (no trazable)":
        "El registro no tiene el PRODUCTOS-ID ENCUESTA. Sin ese identificador no se "
        "puede vincular la encuesta con el marco muestral de beneficiarios ni dar "
        "seguimiento. Se corrige recuperando el ID con el encuestador.",
    "A1a Sin cultivos capturados":
        "El productor declaró tener cultivos, pero el roster de cultivos está "
        "completamente VACÍO: no se registró ninguno (sin área sembrada, producción "
        "ni rendimiento). Es la pérdida más grave del dato productivo, central para "
        "la evaluación agrícola. Revisar por qué no se levantó ningún cultivo.",
    "A1b Roster productivo incompleto":
        "El roster capturó ALGO pero MENOS de lo que el productor declaró: menos "
        "cultivos, o menos parcelas de las declaradas. Falta parte del detalle "
        "productivo. No marca cuando el roster tiene más filas que lo declarado (un "
        "cultivo repartido en varias parcelas genera varias filas y no es un error).",
    "A2 Polígono pendiente":
        "El productor autorizó registrar el polígono de la parcela, pero éste quedó "
        "marcado como PENDIENTE: aún no se ha medido en campo. Es un pendiente "
        "operativo del levantamiento, no un error de dato.",
    "A3 Medida del polígono muy diferente":
        "El área del polígono medida por GPS (en metros cuadrados) es muy diferente "
        "del área que declaró el productor, una vez convertida a metros cuadrados "
        "según su unidad (manzanas, tareas, hectáreas o varas). Se marca cuando la "
        "medida difiere más de 3 veces de la declarada, o cuando la geometría es "
        "degenerada (menos de 3 vértices o área casi nula). Sugiere un error de trazo "
        "del polígono o de la superficie declarada; verificar en campo.",
    "S Composición del hogar incompleta":
        "Falta la composición del hogar (total de personas, mujeres, hombres). Puede "
        "deberse a la lógica de salto del formulario o a una omisión; conviene "
        "verificar contra el XLSForm si el bloque debía desplegarse para este caso.",
    "X Contradicción lógica (SCALL)":
        "Hay respuestas internamente contradictorias: (a) dice usar el agua del SCALL "
        "pero reporta 0 meses de aporte al año; (b) responde que NO tiene un SCALL "
        "instalado, pese a estar en la muestra SCALL (posible problema de elegibilidad "
        "o de captura); o (c) mujeres + hombres no suman el total de personas del "
        "hogar. Revisar el registro para identificar cuál dato es el incorrecto.",
    "X Contradicción lógica (Agrícola)":
        "Hay respuestas internamente contradictorias: (a) edad del productor menor a 18 "
        "o mayor a 90; (b) ingreso por ventas mayor a 10 veces el total de gastos "
        "productivos (rentabilidad implausible, algún monto mal capturado); (c) un "
        "cultivo con área sembrada mayor a cero pero producción cero (posible pérdida "
        "total real o dato faltante); o (d) área cosechada mayor a la sembrada en la "
        "misma unidad. Revisar el registro para verificar el dato.",
}


def _drop9999(s):
    return s.mask(s == 9999).mask(s == 999)


def flags_scall(d: pd.DataFrame, book: dict = None) -> pd.DataFrame:
    f = pd.DataFrame(index=d.index)
    f["Q Sin identificador (no trazable)"] = vacios(col(d, K_ID))
    T = num(col(d, "personas habitan al dia de hoy"))
    M = num(col(d, "cuantas son mujeres"))
    H = num(col(d, "cuantas son hombres"))
    comp = [resolve(d, k) for k in ["personas habitan al dia de hoy",
                                    "cuantas son mujeres", "cuantas son hombres"]]
    comp = [c for c in comp if c is not None]
    if comp:
        f["S Composición del hogar incompleta"] = pd.concat(
            [vacios(d[c]) for c in comp], axis=1).any(axis=1)
    # X — contradicciones lógicas
    usa = col(d, "el hogar usa agua del SCALL").astype(str)
    mes = _drop9999(num(col(d, ["meses al año su scall aporta",
                                "cuantos meses aporta agua"])).mask(lambda s: s == 99))
    inst = col(d, "tiene instalado en su hogar un sistema").astype(str)
    contra = ((usa.str.startswith("Sí") & (mes == 0))
              | inst.str.startswith("No")
              | (T.notna() & M.notna() & H.notna() & ((M + H) != T)))
    f["X Contradicción lógica (SCALL)"] = contra.fillna(False)
    return f.fillna(False)


def encuestas_contradiccion_cultivo(book):
    """Set de llaves de encuesta (_submission__id) con contradicción a nivel
    cultivo: producción=0 con área sembrada>0, o cosechada>sembrada."""
    if book is None:
        return set()
    malos = set()
    for r in hojas_cultivo(book):
        k = clave_encuesta_roster(r)
        if k is None:
            continue
        semb = num(col(r, "area total sembrada"))
        cos = num(col(r, "area total cosechada"))
        prod = num(col(r, "la produccion de"))
        us = col(r, ["M2_Q2b", "M2s_Q2b"], exact=True).astype(str)
        uc = col(r, ["M2_Q3b", "M2s_Q3b"], exact=True).astype(str)
        cond = ((semb > 0) & (prod == 0)) | ((us == uc) & (cos > semb))
        malos |= set(r.loc[cond, k].astype(str).tolist())
    return malos


def contar_roster_por_encuesta(book, sheet):
    """Filas del roster por encuesta, contadas con la llave de encuesta."""
    r = pick_sheet(book, sheet)
    if r is None:
        return pd.Series(dtype=int)
    k = clave_encuesta_roster(r)
    if k is None:
        return pd.Series(dtype=int)
    return r[k].astype(str).value_counts()


def hojas_cultivo(book):
    """Hojas de cultivos ÚNICAS. Según la versión de la exportación puede
    existir 'roster_cultivo', 'roster_cultivos' o ambas; se evita contar dos
    veces la misma hoja (el matcher por subcadena las confunde)."""
    vistas, out = set(), []
    for cand in ("roster_cultivo", "roster_cultivos"):
        r = pick_sheet(book, cand)
        if r is not None and id(r) not in vistas:
            vistas.add(id(r))
            out.append(r)
    return out


def contar_cultivos_por_encuesta(book):
    total = pd.Series(dtype=float)
    for r in hojas_cultivo(book):
        k = clave_encuesta_roster(r)
        if k is None:
            continue
        total = total.add(r[k].astype(str).value_counts(), fill_value=0)
    return total


def parcelas_de_encuesta(book, id_val):
    """Filas de roster_parcela cuya encuesta padre es id_val (usa _submission__id)."""
    rp = pick_sheet(book, "roster_parcela")
    if rp is None:
        return None
    k = clave_encuesta_roster(rp)
    if k is None:
        return rp.iloc[0:0]
    return rp[rp[k].astype(str) == str(id_val)]


def flags_agricola(d: pd.DataFrame, book: dict = None) -> pd.DataFrame:
    f = pd.DataFrame(index=d.index)
    f["Q Sin identificador (no trazable)"] = vacios(col(d, K_ID))

    kc = clave_encuesta_data(d)
    kd = d[kc].astype(str) if kc is not None else id_encuesta(d)
    dec_p = num(col(d, "TERRENOS o PARCELAS"))
    dec_c = num(col(d, "CULTIVOS tuvo en total"))

    # A1a — sin ningún cultivo capturado ; A1b — capturó menos de lo declarado
    a1a = pd.Series(False, index=d.index)
    a1b = pd.Series(False, index=d.index)
    if book is not None:
        pcount = contar_roster_por_encuesta(book, "roster_parcela")
        cpe = contar_cultivos_por_encuesta(book)
        for i in d.index:
            k = kd.iloc[i]
            real_p = pcount.get(k, 0)
            real_c = cpe.get(k, 0)
            dp, dc = dec_p.iloc[i], dec_c.iloc[i]
            # A1a: declaró cultivos pero el roster de cultivos está VACÍO (0).
            sin_cult = (not pd.isna(dc)) and dc > 0 and real_c == 0
            # A1b: capturó ALGO pero menos de lo declarado (cultivos o parcelas).
            #      No marca cuando el roster tiene más (un cultivo en varias
            #      parcelas genera varias filas y no es un error).
            menos = (((not pd.isna(dc)) and 0 < real_c < dc)
                     or ((not pd.isna(dp)) and real_p < dp))
            a1a.iloc[i] = bool(sin_cult)
            a1b.iloc[i] = bool(menos)
    f["A1a Sin cultivos capturados"] = a1a
    f["A1b Roster productivo incompleto"] = a1b

    # A2 — polígono pendiente ; A3 — medida del polígono muy diferente
    geo_c = resolve(d, "coordenadas de la esquina de la parcela")
    a2 = pd.Series(False, index=d.index)   # pendiente
    a3 = pd.Series(False, index=d.index)   # medida muy diferente
    if geo_c is not None:
        rp = pick_sheet(book, "roster_parcela") if book else None
        ac = resolve(rp, "el area de") if rp is not None else None
        uc = resolve(rp, "M1_Q6b", exact=True) if rp is not None else None
        esp_c = resolve(rp, "Especifique", exact=True) if rp is not None else None
        for i in d.index:
            v = d.at[i, geo_c]
            txt = str(v).strip() if v is not None else ""
            tiene_poly = isinstance(v, str) and ";" in str(v)
            # A2: el polígono quedó marcado como PENDIENTE (autorizado, sin medir)
            if txt.upper() == "PENDIENTE":
                a2.iloc[i] = True
                continue
            if not tiene_poly:
                continue
            # A3: comparar área medida (m²) vs área declarada convertida a m²
            pts = parse_geoshape(v)
            am = area_poligono_m2(pts)          # ya en m²
            decl_m2 = np.nan
            if rp is not None and ac is not None:
                mias = parcelas_de_encuesta(book, kd.iloc[i])
                if mias is not None and len(mias):
                    esp = mias[esp_c].iloc[0] if esp_c is not None else None
                    decl_m2, _ = area_declarada_m2(
                        mias[ac].iloc[0],
                        str(mias[uc].iloc[0]) if uc else "", esp)
            muy_dif = (len(pts) < 3 or am < 50
                       or (not pd.isna(decl_m2) and decl_m2 > 0
                           and (am / decl_m2 > 3 or am / decl_m2 < 1/3)))
            a3.iloc[i] = bool(muy_dif)
    f["A2 Polígono pendiente"] = a2
    f["A3 Medida del polígono muy diferente"] = a3

    # X — contradicciones lógicas
    edad = num(col(d, "edad del productor"))
    G = (_drop9999(num(col(d, "compra de la SEMILLA"))).fillna(0)
         + _drop9999(num(col(d, "compra de FERTILIZANTES"))).fillna(0)
         + _drop9999(num(col(d, "AGROQUIMICOS"))).fillna(0)
         + _drop9999(num(col(d, "mano de obra o jornales"))).fillna(0))
    I = _drop9999(num(col(d, "ingreso total obtenido por la venta")))
    contra = (((edad < 18) | (edad > 90))
              | ((I > G * 10) & (G > 0) & I.notna()))
    enc_bad = encuestas_contradiccion_cultivo(book)
    if enc_bad:
        contra = contra | kd.isin(enc_bad)
    f["X Contradicción lógica (Agrícola)"] = contra.fillna(False)
    return f.fillna(False)


def calcular_flags(d, esperado, book):
    return flags_scall(d, book) if esperado == "SCALL" else flags_agricola(d, book)


# ----------------------------------------------------------------------------
# Preguntas clave (valores faltantes)
# ----------------------------------------------------------------------------
CLAVES_SCALL = [
    ("Fecha de la entrevista", "Fecha de la entrevista"),
    ("Nombre del encuestador", K_ENUM),
    ("PRODUCTOS-ID ENCUESTA", K_ID),
    ("Distrito", "Distrito"),
    ("Cantón", "Cantón"),
    ("Geolocalización del hogar", "Registrar geolocalización del hogar"),
    ("Personas en el hogar", "personas habitan al dia de hoy"),
    ("Mujeres en el hogar", "cuantas son mujeres"),
    ("Hombres en el hogar", ["cuantos son hombres", "cuantas son hombres"]),
    ("Año de instalación del SCALL", "año que le instalaron"),
    ("Capacidad de almacenamiento", "capacidad total de almacenamiento"),
]
CLAVES_AGRI = [
    ("Fecha de la entrevista", "Fecha de la entrevista"),
    ("Nombre del encuestador", K_ENUM),
    ("PRODUCTOS-ID ENCUESTA", K_ID),
    ("Distrito", "Distrito"),
    ("Edad del productor", "edad del productor"),
    ("N.º de parcelas", "TERRENOS o PARCELAS"),
    ("N.º de cultivos", "CULTIVOS tuvo en total"),
    ("Gasto en semilla", "compra de la SEMILLA"),
    ("Ingreso por ventas", "ingreso total obtenido por la venta"),
    ("Fuente principal de ingresos", "principal fuente de ingresos"),
]


def faltantes_por_fila(d, claves):
    cols = [resolve(d, k) for _, k in claves]
    cols = [c for c in cols if c is not None]
    if not cols:
        return pd.Series(0, index=d.index)
    return pd.concat([vacios(d[c]) for c in cols], axis=1).sum(axis=1)


# ----------------------------------------------------------------------------
# Componentes visuales
# ----------------------------------------------------------------------------
def hero(titulo, subtitulo):
    st.markdown(
        f'<div class="hero"><h1>{titulo}</h1><p>{subtitulo}</p></div>',
        unsafe_allow_html=True)


def barra(serie, titulo, etiqueta_x, horizontal=False, color=VERDE):
    if horizontal:
        fig = px.bar(serie.sort_values(), orientation="h",
                     labels={"value": "Casos", "index": ""})
    else:
        fig = px.bar(serie, labels={"value": etiqueta_x, "index": ""}, title=titulo)
    fig.update_traces(marker_color=color)
    fig.update_layout(showlegend=False, height=330, title=titulo,
                      plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                      margin=dict(l=0, r=0, t=40, b=0), font=dict(size=12))
    return fig


def kpis(d, flags):
    enum_c = resolve(d, K_ENUM)
    distr_c = resolve(d, "Distrito", exact=True) or resolve(d, "Distrito")
    fechas = pd.to_datetime(col(d, "Fecha de la entrevista"), errors="coerce")
    n_flag = int(flags.any(axis=1).sum())
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Encuestas", len(d))
    c2.metric("Encuestadores", d[enum_c].nunique() if enum_c else "—")
    c3.metric("Distritos", d[distr_c].nunique() if distr_c else "—")
    c4.metric("Días de campo", int(fechas.dt.date.nunique()))
    c5.metric("Con ≥1 bandera",
              f"{n_flag}  ({n_flag/len(d):.0%})" if len(d) else "0")


def resumen_banderas(d, flags):
    st.subheader("🚩 Banderas de calidad")
    resumen = flags.sum()
    resumen = resumen[resumen > 0].sort_values(ascending=False)
    if resumen.empty:
        st.success("Ninguna bandera activa. 🎉")
        return
    t, g = st.columns([5, 4])
    tabla = resumen.rename("Casos").to_frame()
    tabla["% del total"] = (resumen / len(d)).map("{:.0%}".format)
    t.dataframe(tabla, width="stretch")
    g.plotly_chart(barra(resumen, "Casos por bandera", "Casos",
                         horizontal=True, color=ROJO), width="stretch")
    with st.expander("ℹ️ ¿Qué significa cada bandera?"):
        for nombre in resumen.index:
            if nombre in FLAG_DESC:
                st.markdown(f"**{nombre}** — {FLAG_DESC[nombre]}")


def avance_campo(d):
    st.subheader("📈 Avance de campo")
    enum_c = resolve(d, K_ENUM)
    distr_c = resolve(d, "Distrito", exact=True) or resolve(d, "Distrito")
    fechas = pd.to_datetime(col(d, "Fecha de la entrevista"), errors="coerce")
    a, b = st.columns(2)
    if distr_c:
        a.plotly_chart(barra(d[distr_c].value_counts(), "Por distrito", "Encuestas"),
                       width="stretch")
    if enum_c:
        b.plotly_chart(barra(d[enum_c].value_counts(), "Por encuestador", "Encuestas"),
                       width="stretch")
    if fechas.notna().any():
        por_dia = fechas.dt.date.value_counts().sort_index()
        st.plotly_chart(barra(por_dia, "Encuestas por día", "Encuestas", color=VERDE),
                        width="stretch")


def tabla_flags(d, flags):
    st.subheader("⚠ Registros con banderas")
    st.caption("Sin datos personales — use el ID (PRODUCTOS-ID ENCUESTA) para ubicar "
               "el registro en Kobo.")
    base = pd.DataFrame({
        "ID": id_encuesta(d),
        "Encuestador": col(d, K_ENUM),
        "Fecha": pd.to_datetime(col(d, "Fecha de la entrevista"),
                                errors="coerce").dt.strftime("%d/%m/%Y"),
        "Distrito": col(d, "Distrito"),
    })
    if not flags.any(axis=1).any():
        st.info("Sin registros con banderas.")
        return
    sel = st.multiselect("Filtrar por bandera", list(flags.columns),
                         key=f"filtro_{id(flags)}")
    mask = flags[sel].any(axis=1) if sel else flags.any(axis=1)
    out = base[mask].copy()
    out["Banderas"] = flags[mask].apply(
        lambda r: ", ".join(f.split(" ")[0] for f in flags.columns if r[f]), axis=1)
    st.dataframe(pd.concat([out, flags[mask].replace({True: "⚠", False: ""})], axis=1),
                 width="stretch", hide_index=True)


def seccion_faltantes(d, claves):
    st.subheader("🕳 Valores faltantes en preguntas clave")
    st.caption("Parte de los vacíos puede ser lógica de salto del formulario. El valor "
               "está en detectar patrones: una pregunta o un encuestador con faltantes "
               "sistemáticos.")
    filas = []
    for etiqueta, key in claves:
        c = resolve(d, key)
        if c is None:
            filas.append((etiqueta, "—", "—"))
            continue
        v = vacios(d[c])
        filas.append((etiqueta, int(v.sum()), f"{v.mean():.0%}"))
    t1 = pd.DataFrame(filas, columns=["Pregunta clave", "Sin dato", "% del total"])
    a, b = st.columns(2)
    a.markdown("**Por pregunta**")
    a.dataframe(t1, hide_index=True, width="stretch")
    enum_c = resolve(d, K_ENUM)
    if enum_c:
        nf = faltantes_por_fila(d, claves)
        t2 = pd.DataFrame({
            "Encuestas": d.groupby(d[enum_c]).size(),
            "Prom. claves sin dato": nf.groupby(d[enum_c]).mean().round(1),
        }).sort_values("Prom. claves sin dato", ascending=False)
        b.markdown("**Por encuestador**")
        b.dataframe(t2, width="stretch")


# --- Mapas ---
ESTILOS_MAPA = {"⬜ Blanco": "white-bg", "🩶 Claro": "carto-positron",
                "🗺 Calles": "open-street-map"}


def selector_estilo(key):
    return ESTILOS_MAPA[st.radio("Fondo del mapa", list(ESTILOS_MAPA),
                                 horizontal=True, key=key)]


def seccion_mapa(d, modulo=""):
    lat_c, lon_c = resolve(d, "_latitude"), resolve(d, "_longitude")
    if lat_c is None or lon_c is None:
        return
    pts = pd.DataFrame({
        "lat": num(d[lat_c]), "lon": num(d[lon_c]), "ID": id_encuesta(d),
        "Encuestador": col(d, K_ENUM).astype(str), "Distrito": col(d, "Distrito").astype(str),
    }).dropna(subset=["lat", "lon"])
    if pts.empty:
        return
    st.subheader("🗺 Mapa de puntos GPS")
    st.caption("Puntos identificados solo por ID. Uso interno del equipo.")
    fuera = ~(pts["lat"].between(12.9, 14.5) & pts["lon"].between(-90.2, -87.6))
    if fuera.any():
        st.warning("⚠ Punto(s) fuera de El Salvador: IDs "
                   + ", ".join("#" + i for i in pts.loc[fuera, "ID"]))
    estilo = selector_estilo(f"est_pts_{modulo}")
    try:
        fig = px.scatter_map(pts, lat="lat", lon="lon", color="Encuestador",
                             hover_name="ID", hover_data={"lat": False, "lon": False,
                                                          "Distrito": True},
                             zoom=8, height=470)
        fig.update_layout(map_style=estilo, margin=dict(l=0, r=0, t=10, b=0))
    except Exception:
        fig = px.scatter_mapbox(pts, lat="lat", lon="lon", color="Encuestador",
                                hover_name="ID", zoom=8, height=470)
        fig.update_layout(mapbox_style=estilo, margin=dict(l=0, r=0, t=10, b=0))
    fig.update_traces(marker=dict(size=10))
    st.plotly_chart(fig, width="stretch")


def seccion_poligonos(d, book):
    geo_c = resolve(d, "coordenadas de la esquina de la parcela")
    if geo_c is None:
        return
    ids = id_encuesta(d)
    kc = clave_encuesta_data(d)
    kd = d[kc].astype(str) if kc is not None else id_encuesta(d)
    rp = pick_sheet(book, "roster_parcela")
    ac = resolve(rp, "el area de") if rp is not None else None
    uc = resolve(rp, "M1_Q6b", exact=True) if rp is not None else None
    esp_c = resolve(rp, "Especifique", exact=True) if rp is not None else None

    n_pendiente = 0
    regs, shapes = [], []
    for i in d.index:
        v = d.at[i, geo_c]
        if isinstance(v, str) and v.strip().upper() == "PENDIENTE":
            n_pendiente += 1
            continue
        if not (isinstance(v, str) and ";" in str(v)):
            continue
        pts = parse_geoshape(v)
        if not pts:
            continue
        am = area_poligono_m2(pts)                      # m²
        precs = [p[2] for p in pts if not pd.isna(p[2])]
        prec = float(np.mean(precs)) if precs else np.nan
        rep_txt, decl_m2, ratio = "—", np.nan, np.nan
        if rp is not None and ac is not None:
            mias = parcelas_de_encuesta(book, kd.iloc[i])
            if mias is not None and len(mias):
                esp = mias[esp_c].iloc[0] if esp_c is not None else None
                decl_m2, rep_txt = area_declarada_m2(
                    mias[ac].iloc[0], str(mias[uc].iloc[0]) if uc else "", esp)
                if not pd.isna(decl_m2) and decl_m2 > 0:
                    ratio = am / decl_m2
        al = []
        if len(pts) < 3:
            al.append("< 3 vértices")
        if am < 50:
            al.append("área ≈ 0")
        if not pd.isna(ratio) and (ratio > 3 or ratio < 1/3):
            al.append(f"medida muy diferente (x{ratio:.1f})")
        regs.append({
            "ID": ids.loc[i], "Vértices": len(pts),
            "Área medida (m²)": round(am),
            "Área medida (mz)": round(am / M2_POR_MANZANA, 2),
            "Área declarada": rep_txt,
            "Área declarada (m²)": round(decl_m2) if not pd.isna(decl_m2) else "—",
            "⚠ Revisar": "; ".join(al) if al else "✔",
        })
        shapes.append((ids.loc[i], pts))
    if n_pendiente:
        st.info(f"📌 {n_pendiente} parcelas con polígono **PENDIENTE** de medir "
                f"({n_pendiente/len(d):.0%} del total).")
    if not regs:
        return

    st.subheader("📐 Polígonos de parcela (control de calidad)")
    st.caption("Área medida por GPS (en m²) vs. área declarada convertida a m² según "
               "su unidad. Identificados solo por ID. Uso interno del equipo.")
    st.dataframe(pd.DataFrame(regs), hide_index=True, width="stretch")

    import plotly.graph_objects as go
    usa = hasattr(go, "Scattermap")
    estilo = selector_estilo("est_poly")
    fig = go.Figure()
    for pid, pts in shapes:
        la = [p[0] for p in pts] + [pts[0][0]]
        lo = [p[1] for p in pts] + [pts[0][1]]
        tr = dict(lat=la, lon=lo, mode="lines+markers", fill="toself",
                  name=f"#{pid}", hovertext=f"ID #{pid}")
        fig.add_trace(go.Scattermap(**tr) if usa else go.Scattermapbox(**tr))
    centro = dict(lat=float(np.mean([p[0] for _, ps in shapes for p in ps])),
                  lon=float(np.mean([p[1] for _, ps in shapes for p in ps])))
    key = "map" if usa else "mapbox"
    fig.update_layout(**{key: dict(style=estilo, center=centro, zoom=13)},
                      height=480, margin=dict(l=0, r=0, t=10, b=0))
    st.plotly_chart(fig, width="stretch")

    st.markdown("**🔎 Ver el polígono de una observación**")
    opciones = [pid for pid, _ in shapes]
    sel = st.selectbox("Seleccione la encuesta (ID)", opciones, key="poly_sel")
    if sel not in dict(shapes):
        sel = opciones[0]
    pts_sel = dict(shapes)[sel]
    st.dataframe(pd.DataFrame([next((r for r in regs if r["ID"] == sel), regs[0])]),
                 hide_index=True, width="stretch")
    la = [p[0] for p in pts_sel] + [pts_sel[0][0]]
    lo = [p[1] for p in pts_sel] + [pts_sel[0][1]]
    tr = dict(lat=la, lon=lo, mode="lines+markers", fill="toself", name=f"#{sel}")
    fig2 = go.Figure(go.Scattermap(**tr) if usa else go.Scattermapbox(**tr))
    c2 = dict(lat=float(np.mean([p[0] for p in pts_sel])),
              lon=float(np.mean([p[1] for p in pts_sel])))
    fig2.update_layout(**{key: dict(style=estilo, center=c2, zoom=17)},
                       height=440, margin=dict(l=0, r=0, t=10, b=0), showlegend=False)
    st.plotly_chart(fig2, width="stretch")


def poligonos_sin_asignar(book):
    """Lee la hoja 'Poligonos_SINASIGNAR' (polígonos cargados aparte,
    identificados por código de beneficiario). Devuelve [(código, pts), ...]."""
    sp = pick_sheet(book, "Poligonos_SINASIGNAR", "sin asignar", "sinasignar")
    if sp is None or sp.shape[1] < 2:
        return []
    pcol, ccol = sp.columns[0], sp.columns[1]
    out = []
    for _, row in sp.iterrows():
        poly = str(row[pcol])
        cod = str(row[ccol]).strip()
        if ";" not in poly:
            continue
        pts = parse_geoshape(poly)
        if pts:
            out.append((cod if cod and cod != "nan" else "(sin código)", pts))
    return out


def seccion_poligonos_extra(d, book):
    """Aviso de polígonos agregados manualmente (hoja Poligonos_SINASIGNAR) que
    aún no se pueden vincular: su identificador no coincide con ninguna encuesta.
    No se listan ni se mapean; solo se reporta la cantidad."""
    sin = poligonos_sin_asignar(book)
    if not sin:
        return
    st.info(
        f"📍 **{len(sin)} polígonos no asignados aún.** Fueron medidos en el 4.º "
        "enlace (link) de KoboToolbox y cargados aparte. El identificador colocado "
        "a cada uno **no coincide con ningún PRODUCTOS-ID ENCUESTA** de la base, por "
        "lo que todavía no pueden vincularse a un productor ni compararse contra su "
        "área declarada. Pendiente de conciliar los códigos.")


def ficha_encuestado(d, flags, book=None, modulo=""):
    st.subheader("👤 Ficha de la encuesta")
    st.caption("Identificada solo por ID (sin datos personales). Los campos de nombres, "
               "teléfonos, direcciones y coordenadas están excluidos del tablero.")
    ids = id_encuesta(d)
    kobo = _kobo(d)
    etiquetas = ids.where(ids.str.startswith("s/ID"), ids + " · k" + kobo)
    elegido = st.selectbox("Seleccione la encuesta por ID", etiquetas, key=f"ficha_{modulo}")
    m = etiquetas == elegido
    i = etiquetas.index[m][0] if m.any() else etiquetas.index[0]
    activos = [f for f in flags.columns if flags.at[i, f]]
    if activos:
        st.warning("⚠ Banderas de este registro: " + ", ".join(activos))
    else:
        st.success("✔ Sin banderas en este registro.")
    ficha = quitar_sensibles(d).loc[i].dropna()
    ficha = ficha[ficha.astype(str).str.strip() != ""]
    ficha.index.name = "Pregunta"
    st.dataframe(ficha.rename("Respuesta").to_frame().astype(str),
                 width="stretch", height=440)
    if book is not None:
        mias = parcelas_de_encuesta(book, kobo.loc[i])
        if mias is not None and len(mias):
            st.markdown(f"**🌾 Parcelas de esta encuesta ({len(mias)})**")
            st.dataframe(quitar_sensibles(mias).dropna(axis=1, how="all").astype(str),
                         width="stretch", hide_index=True)


# ----------------------------------------------------------------------------
# Exportaciones
# ----------------------------------------------------------------------------
def excel_marcado(d, flags, claves, modulo):
    from openpyxl.styles import PatternFill, Font
    AM = PatternFill("solid", fgColor="FFF3B0")
    RO = PatternFill("solid", fgColor="F5A6A6")
    base = quitar_sensibles(d).copy()
    base.insert(0, "BANDERAS", flags.apply(
        lambda r: ", ".join(f.split(" ")[0] for f in flags.columns if r[f]), axis=1))
    base.insert(0, "ID", id_encuesta(d))
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        base.to_excel(w, index=False, sheet_name="BASE")
        resumen = flags.sum()
        ley = pd.DataFrame(
            [("AMARILLO", "Dato faltante en pregunta clave"),
             ("ROJO", "Registro con bandera de calidad"), ("", "")]
            + [(k, f"{int(v)} caso(s)") for k, v in resumen.items() if v > 0],
            columns=["Marca / Bandera", "Significado / Casos"])
        ley.to_excel(w, index=False, sheet_name="LEYENDA")
        ws = w.book["BASE"]
        ws.freeze_panes = "C2"
        pos = {c: j + 1 for j, c in enumerate(base.columns)}
        for _, key in claves:
            c = resolve(d, key)
            if c is None or c not in pos:
                continue
            vv = vacios(d[c])
            for r_i, idx in enumerate(d.index):
                if vv.loc[idx]:
                    ws.cell(row=r_i + 2, column=pos[c]).fill = AM
        con = flags.any(axis=1)
        for r_i, idx in enumerate(d.index):
            if con.loc[idx]:
                ws.cell(row=r_i + 2, column=pos["BANDERAS"]).fill = RO
        for celda in ws[1]:
            celda.font = Font(bold=True)
        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 20
    return buf.getvalue()


def excel_dummies(d, flags, book, modulo):
    """Devuelve el MISMO libro (todas las hojas y columnas idénticas) y, en la
    hoja principal, agrega una columna dummy 0/1 por bandera: 1 si el registro
    dispara esa bandera, 0 si no. Para revisar las banderas sobre la base cruda."""
    dummies = pd.DataFrame(
        {f"BANDERA_{f.split(' ')[0]}": flags[f].astype(int).to_numpy()
         for f in flags.columns},
        index=range(len(d)))
    main_cols = list(d.columns)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        escritas = set()
        for name, sheet in (book.items() if book else [("DATA", d)]):
            hoja = name[:31]
            if list(sheet.columns) == main_cols and len(sheet) == len(d) and "main" not in escritas:
                out = pd.concat([d.reset_index(drop=True), dummies], axis=1)
                out.to_excel(w, index=False, sheet_name=hoja)
                escritas.add("main")
            else:
                sheet.to_excel(w, index=False, sheet_name=hoja)
        if "main" not in escritas:  # por si no se identificó la hoja principal
            pd.concat([d.reset_index(drop=True), dummies], axis=1).to_excel(
                w, index=False, sheet_name="DATA_con_banderas")
    return buf.getvalue()


def resumen_pdf(d, flags, modulo, extra=""):
    """Resumen de 1 página en PDF (reportlab)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    enum_c = resolve(d, K_ENUM)
    distr_c = resolve(d, "Distrito", exact=True) or resolve(d, "Distrito")
    fechas = pd.to_datetime(col(d, "Fecha de la entrevista"), errors="coerce").dt.date
    con_flag = flags.any(axis=1)
    verde = colors.HexColor("#1F5C2E")
    rojo = colors.HexColor("#D1495B")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=1.3*cm,
                            bottomMargin=1.1*cm, leftMargin=1.5*cm, rightMargin=1.5*cm)
    ss = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=ss["Title"], textColor=verde, fontSize=17, spaceAfter=2)
    sub = ParagraphStyle("s", parent=ss["Normal"], textColor=colors.grey, fontSize=9)
    sec = ParagraphStyle("sec", parent=ss["Heading2"], textColor=verde, fontSize=12,
                         spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("b", parent=ss["Normal"], fontSize=9.5, leading=13)
    el = []
    fv = fechas.dropna()
    periodo = (f"{fv.min():%d/%m/%Y} – {fv.max():%d/%m/%Y}" if len(fv) else "—")
    el.append(Paragraph(f"Resumen de monitoreo — RECLIMA {modulo}", h))
    el.append(Paragraph(f"Periodo: {periodo} &nbsp;·&nbsp; Generado: "
                        f"{datetime.now():%d/%m/%Y %H:%M} &nbsp;·&nbsp; "
                        f"Evaluación final, corredor seco de El Salvador", sub))
    el.append(Spacer(1, 8))

    # KPIs
    kpi = [["Encuestas", "Encuestadores", "Distritos", "Días de campo", "Con ≥1 bandera"],
           [str(len(d)),
            str(d[enum_c].nunique() if enum_c else "—"),
            str(d[distr_c].nunique() if distr_c else "—"),
            str(int(fechas.nunique())),
            f"{int(con_flag.sum())} ({con_flag.mean():.0%})" if len(d) else "0"]]
    tk = Table(kpi, colWidths=[3.5*cm]*5)
    tk.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F1E9")),
        ("TEXTCOLOR", (0, 1), (-1, 1), verde),
        ("FONTSIZE", (0, 0), (-1, 0), 8), ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("TOPPADDING", (0, 1), (-1, 1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 5), ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CDE0D0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white)]))
    el.append(tk)

    # Banderas
    el.append(Paragraph("Banderas de calidad a atender", sec))
    resumen = flags.sum()
    resumen = resumen[resumen > 0].sort_values(ascending=False)
    if resumen.empty:
        el.append(Paragraph("No hay banderas activas.", body))
    else:
        rows = [["Bandera", "Casos", "% del total"]]
        for k, v in resumen.items():
            rows.append([k, str(int(v)), f"{v/len(d):.0%}"])
        tb = Table(rows, colWidths=[11*cm, 2.2*cm, 3*cm])
        tb.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), verde),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("TEXTCOLOR", (1, 1), (1, -1), rojo),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F8F4")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E4DA")),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        el.append(tb)

    # Avance por encuestador
    if enum_c:
        el.append(Paragraph("Avance por encuestador", sec))
        g = d.groupby(d[enum_c]).size().sort_values(ascending=False)
        cf = con_flag.groupby(d[enum_c]).mean()
        rows = [["Encuestador", "Encuestas", "% con bandera"]]
        for name, cnt in g.items():
            rows.append([str(name), str(int(cnt)), f"{cf.get(name, 0):.0%}"])
        tr = Table(rows, colWidths=[8*cm, 3*cm, 5.2*cm])
        tr.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F1E9")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9), ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E4DA")),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
        el.append(tr)

    if extra:
        el.append(Spacer(1, 6))
        el.append(Paragraph(f"<b>Observaciones:</b> {extra}", body))
    el.append(Spacer(1, 8))
    el.append(Paragraph("Documento sin datos personales — los registros se identifican "
                        "por su PRODUCTOS-ID ENCUESTA.", sub))
    doc.build(el)
    return buf.getvalue()


def identificadores(d):
    """DataFrame índice-alineado con INDEX (llave del usuario o Kobo) e
    'Identificador de encuestado' (o 'No disponible' si falta)."""
    kc = clave_encuesta_data(d)
    idx = d[kc].astype(str).str.strip() if kc is not None else _kobo(d)
    id_c = resolve(d, K_ID)
    if id_c is not None:
        raw = d[id_c].astype(str).str.strip()
        ident = raw.mask(vacios(d[id_c]), "No disponible")
    else:
        ident = pd.Series("No disponible", index=d.index)
    return pd.DataFrame({"INDEX": idx, "Identificador de encuestado": ident}, index=d.index)


def analisis_poligonos(d, book):
    """Lista de dicts con la medición y control de calidad de cada polígono de
    parcela capturado en la encuesta. 'desconfiable' = tiene al menos un motivo."""
    geo_c = resolve(d, "coordenadas de la esquina de la parcela")
    if geo_c is None or book is None:
        return []
    ident = identificadores(d)
    kc = clave_encuesta_data(d)
    kd = d[kc].astype(str) if kc is not None else id_encuesta(d)
    rp = pick_sheet(book, "roster_parcela")
    ac = resolve(rp, "el area de") if rp is not None else None
    uc = resolve(rp, "M1_Q6b", exact=True) if rp is not None else None
    esp_c = resolve(rp, "Especifique", exact=True) if rp is not None else None
    out = []
    for i in d.index:
        v = d.at[i, geo_c]
        if not (isinstance(v, str) and ";" in str(v)):
            continue
        pts = parse_geoshape(v)
        if not pts:
            continue
        am = area_poligono_m2(pts)                       # m²
        rep_txt, decl_m2, ratio = "—", np.nan, np.nan
        if rp is not None and ac is not None:
            mias = parcelas_de_encuesta(book, kd.iloc[i])
            if mias is not None and len(mias):
                esp = mias[esp_c].iloc[0] if esp_c is not None else None
                decl_m2, rep_txt = area_declarada_m2(
                    mias[ac].iloc[0], str(mias[uc].iloc[0]) if uc else "", esp)
                if not pd.isna(decl_m2) and decl_m2 > 0:
                    ratio = am / decl_m2
        motivos = []
        if len(pts) < 3:
            motivos.append("menos de 3 vértices")
        if am < 50:
            motivos.append("área casi nula (≈0)")
        if not pd.isna(ratio) and (ratio > 3 or ratio < 1/3):
            motivos.append(f"medida muy diferente de la declarada (x{ratio:.1f})")
        out.append({
            "INDEX": ident.at[i, "INDEX"],
            "Identificador": ident.at[i, "Identificador de encuestado"],
            "Vertices": len(pts),
            "Area_medida_m2": round(am),
            "Area_medida_mz": round(am / M2_POR_MANZANA, 2),
            "Area_declarada": rep_txt,
            "Area_declarada_m2": round(decl_m2) if not pd.isna(decl_m2) else "—",
            "motivos": motivos,
            "desconfiable": bool(motivos),
        })
    return out


def reporte_docx(d, flags, modulo, extra="", book=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    VERDE = RGBColor(0x1F, 0x5C, 0x2E)

    enum_c = resolve(d, K_ENUM)
    distr_c = resolve(d, "Distrito", exact=True) or resolve(d, "Distrito")
    ident = identificadores(d)
    fechas = pd.to_datetime(col(d, "Fecha de la entrevista"), errors="coerce").dt.date
    con = flags.any(axis=1)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    def _tab(headers, filas, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for j, h in enumerate(headers):
            r = t.rows[0].cells[j].paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(8.5)
        for fila in filas:
            cs = t.add_row().cells
            for j, val in enumerate(fila):
                r = cs[j].paragraphs[0].add_run("" if val is None else str(val))
                r.font.size = Pt(8.5)
        return t

    doc.add_heading(f"Reporte de análisis de campo — RECLIMA {modulo}", 0)
    fv = fechas.dropna()
    rango = f"Periodo: {fv.min():%d/%m/%Y} – {fv.max():%d/%m/%Y}   ·   " if len(fv) else ""
    doc.add_paragraph(rango + f"Generado: {datetime.now():%d/%m/%Y %H:%M}. "
                      "Documento sin datos personales; los registros se identifican por "
                      "su INDEX y su PRODUCTOS-ID ENCUESTA.")

    # 1. Resumen de campo
    doc.add_heading("1. Resumen de campo", 1)
    _tab(["Indicador", "Valor"], [
        ["Encuestas", len(d)],
        ["Encuestadores", d[enum_c].nunique() if enum_c else "—"],
        ["Distritos", d[distr_c].nunique() if distr_c else "—"],
        ["Días de campo", int(fechas.nunique())],
        ["Encuestas con ≥1 bandera", f"{int(con.sum())} ({con.mean():.0%})" if len(d) else "0"],
    ])

    # 2. Estadísticas generales
    doc.add_heading("2. Estadísticas generales (preliminares)", 1)
    doc.add_paragraph("Cálculos sobre la base sin depuración final; referencia de avance, "
                      "no resultados de la evaluación.")
    _tab(["Estadística", "Valor"], [[l, v] for (l, v, *_) in stats_datos(d, modulo)])

    # 3. Avance por día
    doc.add_heading("3. Avance por día", 1)
    por_dia = fechas.dropna().value_counts().sort_index()
    _tab(["Día", "Encuestas"], [[f"{k:%d/%m/%Y}", int(v)] for k, v in por_dia.items()])

    # 4. Banderas: total y significado
    doc.add_heading("4. Banderas de calidad", 1)
    resumen = flags.sum()
    resumen = resumen[resumen > 0].sort_values(ascending=False)
    if resumen.empty:
        doc.add_paragraph("No hay banderas activas.")
    else:
        _tab(["Bandera", "Casos", "% del total"],
             [[k, int(v), f"{v/len(d):.0%}"] for k, v in resumen.items()])
        doc.add_paragraph("Qué significa cada bandera activa:", style="Intense Quote")
        for k in resumen.index:
            if k in FLAG_DESC:
                p = doc.add_paragraph()
                p.add_run(f"{k}: ").bold = True
                p.add_run(FLAG_DESC[k])

    # 5. Registros por bandera (INDEX + identificador)
    doc.add_heading("5. Registros de los que se desconfía, por bandera", 1)
    if resumen.empty:
        doc.add_paragraph("No hay registros con banderas.")
    else:
        for k in resumen.index:
            doc.add_paragraph(f"{k} — {int(resumen[k])} caso(s)", style="Intense Quote")
            filas = [[ident.at[i, "INDEX"], ident.at[i, "Identificador de encuestado"],
                      d[enum_c].iloc[i] if enum_c else "—"]
                     for i in d.index[flags[k]]]
            _tab(["INDEX", "Identificador de encuestado", "Encuestador"], filas)

    # 6. Polígonos desconfiables (solo Agrícola)
    polys = analisis_poligonos(d, book) if modulo != "SCALL" else []
    geo_c0 = resolve(d, "coordenadas de la esquina de la parcela")
    n_pend = int(d[geo_c0].astype(str).str.strip().str.upper().eq("PENDIENTE").sum()) if geo_c0 is not None else 0
    if polys or n_pend:
        doc.add_heading("6. Polígonos de parcela", 1)
        if n_pend:
            doc.add_paragraph(
                f"{n_pend} parcelas tienen el polígono marcado como PENDIENTE de medir "
                f"({n_pend/len(d):.0%} del total): el productor autorizó el registro pero "
                "aún no se ha levantado.")
        malos = [p for p in polys if p["desconfiable"]]
        doc.add_paragraph(
            f"Se capturaron {len(polys)} polígonos. El área medida por GPS se calcula en "
            "metros cuadrados y se compara con el área declarada por el productor, "
            "convertida a metros cuadrados según su unidad (manzanas, tareas, hectáreas o "
            f"varas). De ellos, {len(malos)} resultan desconfiables por al menos uno de "
            "estos criterios: menos de 3 vértices; área casi nula; o medida que difiere "
            "más de 3 veces de la declarada.")
        if malos:
            _tab(["INDEX", "Identificador", "Vért.", "Medida (m²)", "Declarada (m²)",
                  "Declarada", "Motivo"],
                 [[p["INDEX"], p["Identificador"], p["Vertices"], p["Area_medida_m2"],
                   p["Area_declarada_m2"], p["Area_declarada"], "; ".join(p["motivos"])]
                  for p in malos])
        # nota de polígonos sin asignar
        sin = poligonos_sin_asignar(book) if book else []
        if sin:
            doc.add_paragraph(
                f"Además, {len(sin)} polígonos fueron cargados aparte (medidos en el 4.º "
                "enlace de KoboToolbox) con un identificador que no coincide con ningún "
                "PRODUCTOS-ID ENCUESTA de la base, por lo que aún no pueden vincularse.")

    if extra:
        doc.add_heading("Observaciones", 1)
        doc.add_paragraph(extra)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def seccion_reporte(d, flags, modulo, book=None):
    st.subheader("📤 Exportar")
    extra = st.text_input("Observaciones (opcional, se incluyen en los reportes)",
                          key=f"obs_{modulo}")
    claves = CLAVES_SCALL if modulo == "SCALL" else CLAVES_AGRI
    hoy = f"{datetime.now():%Y-%m-%d}"
    c1, c2, c3 = st.columns(3)
    c1.download_button(
        "📄 Resumen 1 página (.pdf)", data=resumen_pdf(d, flags, modulo, extra),
        file_name=f"Resumen_RECLIMA_{modulo}_{hoy}.pdf", mime="application/pdf",
        key=f"pdf_{modulo}", width="stretch")
    c2.download_button(
        "📊 Base marcada (.xlsx)", data=excel_marcado(d, flags, claves, modulo),
        file_name=f"Base_marcada_{modulo}_{hoy}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"xl_{modulo}", width="stretch")
    c3.download_button(
        "📝 Reporte (.docx)", data=reporte_docx(d, flags, modulo, extra, book=book),
        file_name=f"Reporte_RECLIMA_{modulo}_{hoy}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"doc_{modulo}", width="stretch")
    st.download_button(
        "🧮 Base idéntica + columnas dummy por bandera (.xlsx)",
        data=excel_dummies(d, flags, book, modulo),
        file_name=f"Base_dummies_{modulo}_{hoy}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"dummy_{modulo}", width="stretch")
    st.caption("El PDF es un resumen ejecutivo; la base marcada tiene faltantes en "
               "amarillo y banderas en rojo (sin datos personales); el Word detalla "
               "los registros por bandera. La **base idéntica + dummies** es el mismo "
               "archivo del hosting, con las mismas columnas, más una columna 0/1 por "
               "bandera (1 = aplica) para revisarlas sobre la base cruda — incluye "
               "todos los campos, es para revisión interna del administrador.")


# ----------------------------------------------------------------------------
# Estadísticas (preliminares) — por módulo y cruce territorial
# ----------------------------------------------------------------------------
def _pct(serie, cond):
    v = serie.dropna()
    if v.empty:
        return "—"
    return f"{cond(v.astype(str)).mean():.0%}"


def tarjetas(items, por_fila=4):
    for k in range(0, len(items), por_fila):
        cols = st.columns(por_fila)
        for c, (lab, val, *ayuda) in zip(cols, items[k:k + por_fila]):
            c.metric(lab, val, help=ayuda[0] if ayuda else None)


def stats_datos(d, modulo):
    """Lista de (etiqueta, valor) de estadísticas preliminares del módulo.
    Se usa tanto en el tablero como en el reporte Word."""
    if modulo == "SCALL":
        usa = col(d, "el hogar usa agua del SCALL")
        mes = num(col(d, ["meses al año su scall aporta", "cuantos meses aporta agua"])).mask(lambda s: s == 99)
        disp = col(d, "la disponibilidad de agua para beber es")
        seg = col(d, "seguira funcionando dentro de 2")
        fies = [c for c in d.columns if "por falta de dinero" in _norm(c)
                or "ultimos tres meses" in _norm(c)]
        inseg = (d[fies].apply(lambda r: (r.astype(str) == "Sí").any(), axis=1)
                 if fies else pd.Series(dtype=bool))
        return [
            ("Hogares que usan el agua del SCALL", _pct(usa, lambda v: v.str.startswith("Sí"))),
            ("Meses/año de aporte (prom.)", f"{mes.mean():.1f}" if mes.notna().any() else "—"),
            ("Perciben mejor disponibilidad de agua", _pct(disp, lambda v: v.isin(["Mejor", "Mucho mejor"]))),
            ("Creen que funcionará en 2 años", _pct(seg, lambda v: v.str.startswith("Sí"))),
            ("Con algún signo de inseguridad alimentaria (FIES)",
             f"{inseg.mean():.0%}" if len(inseg) else "—",
             "Al menos una respuesta afirmativa en la escala FIES."),
        ]
    sx = col(d, "sexo del productor")
    ed = num(col(d, "edad del productor"))
    hg = num(col(d, "personas habitan al dia de hoy"))
    pa = num(col(d, "TERRENOS o PARCELAS"))
    cu = num(col(d, "CULTIVOS tuvo en total"))
    ing = num(col(d, "ingreso total obtenido por la venta")).mask(lambda s: s == 9999)
    gs = num(col(d, "compra de la SEMILLA")).mask(lambda s: s == 9999)
    fu = col(d, "principal fuente de ingresos").dropna().astype(str)
    eca = col(d, "Escuela de Campo")
    pf = col(d, "planes de finca")
    items = [
        ("Productoras mujeres", _pct(sx, lambda v: v == "Mujer")),
        ("Edad promedio del productor", f"{ed.mean():.0f} años" if ed.notna().any() else "—"),
        ("Tamaño promedio del hogar", f"{hg.mean():.1f}" if hg.notna().any() else "—"),
        ("Parcelas por productor (prom.)", f"{pa.mean():.1f}" if pa.notna().any() else "—"),
        ("Cultivos por productor (prom.)", f"{cu.mean():.1f}" if cu.notna().any() else "—"),
        ("Vendieron parte de su cosecha",
         f"{(ing > 0).sum()/ing.notna().sum():.0%}" if ing.notna().any() else "—"),
        ("Ingreso por ventas (prom., >0)", f"${ing[ing > 0].mean():.0f}" if (ing > 0).any() else "—"),
        ("Gasto en semilla (prom., >0)", f"${gs[gs > 0].mean():.0f}" if (gs > 0).any() else "—"),
        ("Participaron en Escuela de Campo", _pct(eca, lambda v: v.str.startswith("Sí"))),
        ("Participaron en planes de finca", _pct(pf, lambda v: v.str.startswith("Sí"))),
    ]
    if len(fu):
        items.append(("Fuente principal de ingresos (moda)",
                      f"{fu.value_counts().index[0]} ({fu.value_counts().iloc[0]/len(fu):.0%})"))
    return items


def stats_scall(d):
    st.markdown("#### 💧 SCALL")
    tarjetas(stats_datos(d, "SCALL"), por_fila=5)


def stats_agricola(d):
    st.markdown("#### 🌾 Prácticas agrícolas")
    items = stats_datos(d, "AGRICOLA")
    moda = next((v for l, v, *_ in items if "Fuente principal" in l), None)
    tarjetas([it for it in items if "Fuente principal" not in it[0]], por_fila=5)
    if moda:
        st.caption(f"Fuente principal de ingresos más común: **{moda}**.")


def cruce_territorial(ds, da):
    st.markdown("#### 🔗 Cruce territorial de las dos bases")
    st.caption("Las encuestas SCALL y Prácticas son de **beneficiarios distintos** "
               "(componentes diferentes del proyecto), por lo que no se pueden cruzar "
               "a nivel de persona. Sí se comparan por territorio: dónde coincide la "
               "cobertura de ambas intervenciones.")
    ms = resolve(ds, "Municipio")
    ma = resolve(da, "Municipio")
    if ms is None or ma is None:
        st.info("No se encontró la columna de municipio en alguna base.")
        return
    cs = ds[ms].dropna().astype(str).value_counts()
    ca = da[ma].dropna().astype(str).value_counts()
    muni = sorted(set(cs.index) | set(ca.index))
    tabla = pd.DataFrame({
        "Municipio": muni,
        "Encuestas SCALL": [int(cs.get(m, 0)) for m in muni],
        "Encuestas Prácticas": [int(ca.get(m, 0)) for m in muni],
    })
    tabla["Ambas intervenciones"] = np.where(
        (tabla["Encuestas SCALL"] > 0) & (tabla["Encuestas Prácticas"] > 0), "✔", "")
    ambos = int((tabla["Ambas intervenciones"] == "✔").sum())
    c1, c2, c3 = st.columns(3)
    c1.metric("Municipios con SCALL", int((tabla["Encuestas SCALL"] > 0).sum()))
    c2.metric("Municipios con Prácticas", int((tabla["Encuestas Prácticas"] > 0).sum()))
    c3.metric("Municipios con ambas", ambos)
    st.dataframe(tabla.sort_values(["Ambas intervenciones", "Encuestas SCALL"],
                                   ascending=[False, False]),
                 hide_index=True, width="stretch")


def pestana_estadisticas():
    st.subheader("📊 Estadísticas preliminares")
    st.caption("Calculadas sobre las bases publicadas, sin depuración final ni "
               "ponderación. Son referencia de avance, **no** resultados de la "
               "evaluación.")
    libros = {}
    for mod_ in ("SCALL", "AGRICOLA"):
        p = os.path.join(DATA_DIR, ARCHIVO[mod_])
        if os.path.exists(p):
            _, d = detect_module(libro_publicado(ARCHIVO[mod_], os.path.getmtime(p)))
            libros[mod_] = d
    if "SCALL" in libros:
        stats_scall(libros["SCALL"])
        st.divider()
    if "AGRICOLA" in libros:
        stats_agricola(libros["AGRICOLA"])
        st.divider()
    if "SCALL" in libros and "AGRICOLA" in libros:
        cruce_territorial(libros["SCALL"], libros["AGRICOLA"])
    elif not libros:
        st.info("Aún no hay bases publicadas. Sube data/scall.xlsx y data/agricola.xlsx.")


# ----------------------------------------------------------------------------
# Render de un módulo
# ----------------------------------------------------------------------------
def _seccion(nombre_seccion, fn, *args):
    """Ejecuta una sección aislando errores: si una falla, muestra un aviso
    con el detalle y el resto del tablero sigue funcionando."""
    try:
        fn(*args)
    except Exception as e:
        import traceback
        st.warning(f"⚠ No se pudo mostrar la sección «{nombre_seccion}»: {e}")
        with st.expander("Ver detalle técnico"):
            st.code(traceback.format_exc())


def render_modulo(book, esperado, nombre):
    detectado, d = detect_module(book)
    if detectado and detectado != esperado:
        st.error(f"Este archivo parece del módulo **{detectado}**, no de {nombre}. "
                 "Súbelo en la otra pestaña.")
        return
    # Índice de fila único (las bases con varias versiones traen _index repetido)
    d = d.reset_index(drop=True)
    try:
        flags = calcular_flags(d, esperado, book)
    except Exception as e:
        import traceback
        st.error(f"No se pudieron calcular las banderas: {e}")
        st.code(traceback.format_exc())
        flags = pd.DataFrame(index=d.index)
    claves = CLAVES_SCALL if esperado == "SCALL" else CLAVES_AGRI
    mod_ = "SCALL" if esperado == "SCALL" else "Agrícola"
    _seccion("Indicadores", kpis, d, flags)
    st.divider()
    _seccion("Resumen de banderas", resumen_banderas, d, flags)
    st.divider()
    _seccion("Avance de campo", avance_campo, d)
    st.divider()
    _seccion("Registros con banderas", tabla_flags, d, flags)
    _seccion("Valores faltantes", seccion_faltantes, d, claves)
    if esperado == "AGRICOLA":
        st.divider()
        _seccion("Polígonos de parcela", seccion_poligonos, d, book)
        _seccion("Polígonos agregados manualmente", seccion_poligonos_extra, d, book)
    st.divider()
    _seccion("Mapa de puntos GPS", seccion_mapa, d, esperado)
    st.divider()
    _seccion("Exportar", seccion_reporte, d, flags, mod_, book)
    st.divider()
    _seccion("Ficha de la encuesta", ficha_encuestado, d, flags,
             book if esperado == "AGRICOLA" else None, esperado)


# ----------------------------------------------------------------------------
# Datos publicados por el administrador
# ----------------------------------------------------------------------------
DATA_DIR = "data"
ARCHIVO = {"SCALL": "scall.xlsx", "AGRICOLA": "agricola.xlsx"}


@st.cache_data(show_spinner=False)
def libro_publicado(nombre_archivo, mtime):
    with open(os.path.join(DATA_DIR, nombre_archivo), "rb") as fh:
        return load_book(fh.read())


def pestana_modulo(esperado, nombre, key):
    publicado = os.path.join(DATA_DIR, ARCHIVO[esperado])
    with st.expander("🔄 Ver otra base (opcional — no reemplaza la publicada)"):
        up = st.file_uploader(f"Base {nombre} (.xlsx)", type="xlsx", key=f"up_{key}")
    if up:
        st.caption("Mostrando la base subida en esta sesión (no queda guardada).")
        render_modulo(load_book(up.getvalue()), esperado, nombre)
    elif os.path.exists(publicado):
        mt = os.path.getmtime(publicado)
        st.caption(f"📌 Base publicada por el administrador — actualizada el "
                   f"{datetime.fromtimestamp(mt):%d/%m/%Y %H:%M}.")
        render_modulo(libro_publicado(ARCHIVO[esperado], mt), esperado, nombre)
    else:
        st.info(f"Aún no hay base publicada para {nombre}. Sube `data/{ARCHIVO[esperado]}` "
                "al repo, o usa la sección de arriba para verla en esta sesión.")


# ----------------------------------------------------------------------------
# App
# ----------------------------------------------------------------------------
if check_password():
    st.markdown(CSS, unsafe_allow_html=True)
    hero("Monitoreo de calidad — RECLIMA",
         "Evaluación final · corredor seco de El Salvador. Datos publicados por el "
         "administrador, sin información personal de los entrevistados.")
    tab_scall, tab_agri, tab_stats = st.tabs(
        ["💧 SCALL", "🌾 Agrícola", "📊 Estadísticas"])
    with tab_scall:
        pestana_modulo("SCALL", "SCALL", "scall")
    with tab_agri:
        pestana_modulo("AGRICOLA", "Agrícola", "agri")
    with tab_stats:
        pestana_estadisticas()
